from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
from datetime import datetime
import json

JSON_FILE = "recommendations.json"
OUTPUT_DIR = "recommendations"

def load_recommendations():
    """Загружает рекомендации из JSON-файла"""
    if not os.path.exists(JSON_FILE):
        return {}
    with open(JSON_FILE, "r", encoding="utf-8") as f:
        return json.load(f)

def generate_word(patient_name, diseases, recommendations):
    """
    Генерирует Word-документ с рекомендациями.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    # Формируем имя файла
    date_str = datetime.now().strftime("%Y-%m-%d_%H-%M")
    file_name = f"{patient_name}_{date_str}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)

    # Загружаем рекомендации
    all_recommendations = load_recommendations()
    general_recommendation = all_recommendations.get("Общие рекомендации", "Следуйте советам врача и ведите здоровый образ жизни.")

    # Создаём документ
    doc = Document()

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run("Медицинские рекомендации")
    title_run.bold = True
    title_run.font.size = Pt(16)

    # Врач
    doc.add_paragraph("Врач: Хвостунова Татьяна Викторовна").runs[0].font.size = Pt(12)

    # Пациент
    patient_paragraph = doc.add_paragraph()
    patient_paragraph.add_run(f"Пациент: {patient_name}").bold = True
    patient_paragraph.runs[0].font.size = Pt(14)

    # Общие рекомендации
    doc.add_paragraph().add_run("Общие рекомендации:").bold = True
    general_paragraph = doc.add_paragraph(general_recommendation)
    general_paragraph.runs[0].font.size = Pt(12)

    doc.add_paragraph()  # Отступ

    # Назначенные рекомендации
    doc.add_paragraph().add_run("Назначенные рекомендации:").bold = True

    for disease in diseases:
        recommendation = recommendations.get(disease, "Нет данных")

        # Заголовок болезни (жирный)
        disease_paragraph = doc.add_paragraph()
        disease_paragraph.add_run(f"{disease}:").bold = True
        disease_paragraph.runs[0].font.size = Pt(12)

        # Рекомендация
        rec_paragraph = doc.add_paragraph(recommendation)
        rec_paragraph.runs[0].font.size = Pt(12)

        doc.add_paragraph()  # Отступ между рекомендациями

    # Сохранение документа
    doc.save(file_path)
    return file_path
